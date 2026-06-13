from __future__ import annotations

import csv
import io
import json
import re
import shutil
import stat
import subprocess
import sys
import tarfile
import tempfile
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Callable

from source.solution.mcp.api_test_executor import execute_api_test_plan
from source.solution.mcp.evidence_chain import analyze_evidence_chain


def _safe_archive_target(output_root: Path, member_name: str) -> Path:
    normalized = member_name.replace("\\", "/")
    relative = Path(normalized)
    if relative.is_absolute() or ".." in relative.parts:
        raise ValueError(f"Unsafe archive member: {member_name}")
    target = (output_root / relative).resolve()
    if target == output_root or not target.is_relative_to(output_root):
        raise ValueError(f"Unsafe archive member: {member_name}")
    return target


def _java_main_class_name(code: str) -> str | None:
    public_match = re.search(
        r"\bpublic\s+(?:(?:final|abstract|strictfp)\s+)*class\s+([A-Za-z_$][\w$]*)",
        code,
    )
    if public_match:
        return public_match.group(1)

    for class_match in re.finditer(
        r"\bclass\s+([A-Za-z_$][\w$]*)",
        code,
    ):
        class_body = code[class_match.end():]
        if re.search(r"\bstatic\s+(?:public\s+)?void\s+main\s*\(", class_body):
            return class_match.group(1)
        if re.search(r"\bpublic\s+static\s+void\s+main\s*\(", class_body):
            return class_match.group(1)
    return None


_DATE_FULL_PATTERN = re.compile(
    r"(\d{4})\s*(?:年|[./-])\s*(\d{1,2})\s*(?:月|[./-])\s*(\d{1,2})"
)
_DATE_NUMBER_PATTERN = r"(\d+|[零〇一二两三四五六七八九十百]+)"
_DATE_WEEKDAY_NUMBERS = {
    "一": 0,
    "二": 1,
    "三": 2,
    "四": 3,
    "五": 4,
    "六": 5,
    "日": 6,
    "天": 6,
}
_DATE_WEEKDAYS = [
    "Monday",
    "Tuesday",
    "Wednesday",
    "Thursday",
    "Friday",
    "Saturday",
    "Sunday",
]


def _chinese_number_to_int(text: str) -> int:
    if text.isdigit():
        return int(text)

    digits = {
        "零": 0,
        "〇": 0,
        "一": 1,
        "二": 2,
        "两": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
    }
    total = 0
    current = 0
    for char in text:
        if char in digits:
            current = digits[char]
        elif char == "十":
            total += (current or 1) * 10
            current = 0
        elif char == "百":
            total += (current or 1) * 100
            current = 0
    return total + current


def _parse_date_datetime(expression: str, year: int, month: int, day: int, start: int = 0) -> datetime:
    tail = expression[start:]
    hour_match = re.match(r"\s*(?:日)?\s*(\d{1,2})(?:点|时|[:.])(?:(\d{1,2}))?", tail)
    hour = int(hour_match.group(1)) if hour_match else 0
    minute = int(hour_match.group(2)) if hour_match and hour_match.group(2) else 0
    return datetime(year, month, day, hour, minute)


def _parse_date_base(expression: str, base_date: str = "") -> tuple[datetime, list[re.Match[str]]]:
    explicit_dates = list(_DATE_FULL_PATTERN.finditer(expression))
    if base_date:
        try:
            return datetime.fromisoformat(base_date.strip().replace(".", "-").replace("/", "-")), explicit_dates
        except ValueError:
            base_match = _DATE_FULL_PATTERN.search(base_date)
            if not base_match:
                raise
            return (
                datetime(
                    int(base_match.group(1)),
                    int(base_match.group(2)),
                    int(base_match.group(3)),
                ),
                explicit_dates,
            )

    if explicit_dates:
        match = explicit_dates[0]
        return (
            _parse_date_datetime(
                expression,
                int(match.group(1)),
                int(match.group(2)),
                int(match.group(3)),
                match.end(),
            ),
            explicit_dates,
        )

    european = re.search(r"(?<!\d)(\d{1,2})/(\d{1,2})/(\d{4})(?!\d)", expression)
    month_day = re.search(r"(\d{1,2})月(\d{1,2})日?", expression)
    year = re.search(r"(20\d{2})年", expression)
    if european:
        return (
            datetime(
                int(european.group(3)),
                int(european.group(2)),
                int(european.group(1)),
            ),
            explicit_dates,
        )
    if month_day and year:
        return (
            datetime(
                int(year.group(1)),
                int(month_day.group(1)),
                int(month_day.group(2)),
            ),
            explicit_dates,
        )
    return datetime.now(), explicit_dates


def _add_months(base: datetime, months: int) -> datetime:
    import calendar

    month_index = base.month - 1 + months
    year = base.year + month_index // 12
    month = month_index % 12 + 1
    day = min(base.day, calendar.monthrange(year, month)[1])
    return base.replace(year=year, month=month, day=day)


def _add_workdays(base: datetime, days: int) -> datetime:
    current = base
    count = 0
    step = 1 if days >= 0 else -1
    while count < abs(days):
        current += timedelta(days=step)
        if current.weekday() < 5:
            count += 1
    return current


def _date_result_payload(expression: str, base: datetime, result: datetime) -> dict[str, Any]:
    return {
        "expression": expression,
        "base_date": base.strftime("%Y-%m-%d"),
        "result": result.strftime("%Y-%m-%d"),
        "weekday": _DATE_WEEKDAYS[result.weekday()],
    }


def _compute_date_payload(expression: str, base_date: str = "") -> dict[str, Any]:
    """Compute one natural-language date expression with deterministic rules."""
    base, explicit_dates = _parse_date_base(expression, base_date)
    exp_lower = expression.lower()

    if not base_date:
        anchor = re.search(
            r"下周([一二三四五六日天])是\s*" + _DATE_FULL_PATTERN.pattern,
            expression,
        )
        if anchor:
            known_weekday = _DATE_WEEKDAY_NUMBERS[anchor.group(1)]
            known_date = datetime(
                int(anchor.group(2)),
                int(anchor.group(3)),
                int(anchor.group(4)),
            )
            base = known_date - timedelta(days=7 + known_weekday)

    result = None
    fiscal_week = re.search(
        r"财年第\s*(\d+)\s*周从\s*"
        + _DATE_FULL_PATTERN.pattern
        + r".*?第\s*\1\s*周的周([一二三四五六日天])",
        expression,
    )
    relative_week_matches = list(
        re.finditer(
            r"(上|下|本|这)(?:一?个|个)?(?:周|星期)(?:周)?([一二三四五六日天])",
            expression,
        )
    )
    relative_week = relative_week_matches[-1] if relative_week_matches else None
    named_weekday = re.search(r"第\s*\d+\s*周的周([一二三四五六日天])", expression)
    month_day_target = re.fullmatch(r"\s*(\d{1,2})月(\d{1,2})日?\s*", expression)

    backward_days = re.search(
        rf"(?:往前推|向前推|倒推|提前)\s*{_DATE_NUMBER_PATTERN}\s*天",
        expression,
    ) or re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*天\s*(?:前|之前|以前)",
        expression,
    ) or re.search(
        rf"(?:前|之前|以前)\s*{_DATE_NUMBER_PATTERN}\s*天",
        expression,
    )
    forward_days = re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*个?\s*(?:自然日|天)\s*(?:后|之后|以后)",
        expression,
    ) or re.search(
        rf"(?:后|之后|以后)\s*{_DATE_NUMBER_PATTERN}\s*个?\s*(?:自然日|天)",
        expression,
    )
    natural_days = re.search(rf"{_DATE_NUMBER_PATTERN}\s*个?\s*自然日", expression)
    hour_offset = re.search(rf"{_DATE_NUMBER_PATTERN}\s*个?\s*小时", expression)
    backward_weeks = re.search(
        rf"(?:往前推|向前推|倒推|提前)\s*{_DATE_NUMBER_PATTERN}\s*(?:周|星期)",
        expression,
    ) or re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*(?:周|星期)\s*(?:前|之前|以前)",
        expression,
    )
    forward_weeks = re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*(?:周|星期)\s*(?:后|之后|以后)",
        expression,
    )
    workday_offset = re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*个?\s*工作日\s*(?:后|之后|以后)",
        expression,
    ) or re.search(
        rf"(?:后|之后|以后)\s*{_DATE_NUMBER_PATTERN}\s*个?\s*工作日",
        expression,
    )
    backward_workday_offset = re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*个?\s*工作日\s*(?:前|之前|以前)",
        expression,
    ) or re.search(
        rf"(?:往前推|向前推|倒推|提前)\s*{_DATE_NUMBER_PATTERN}\s*个?\s*工作日",
        expression,
    )
    month_offset = re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*个?月\s*(?:后|之后|以后)",
        expression,
    )
    backward_month_offset = re.search(
        rf"{_DATE_NUMBER_PATTERN}\s*个?月\s*(?:前|之前|以前)",
        expression,
    )

    if fiscal_week:
        week_start = datetime(
            int(fiscal_week.group(2)),
            int(fiscal_week.group(3)),
            int(fiscal_week.group(4)),
        )
        result = week_start + timedelta(
            days=_DATE_WEEKDAY_NUMBERS[fiscal_week.group(5)] - week_start.weekday()
        )
    elif named_weekday:
        result = base + timedelta(
            days=_DATE_WEEKDAY_NUMBERS[named_weekday.group(1)] - base.weekday()
        )
    elif "大后天" in expression:
        result = base + timedelta(days=3)
    elif "后天" in expression:
        result = base + timedelta(days=2)
    elif "明天" in expression or "明日" in expression:
        result = base + timedelta(days=1)
    elif "前天" in expression:
        result = base - timedelta(days=2)
    elif "昨天" in expression or "昨日" in expression:
        result = base - timedelta(days=1)
    elif relative_week:
        week_offset = {"上": -1, "本": 0, "这": 0, "下": 1}[relative_week.group(1)]
        week_start = base - timedelta(days=base.weekday())
        result = week_start + timedelta(
            weeks=week_offset,
            days=_DATE_WEEKDAY_NUMBERS[relative_week.group(2)],
        )
    elif "last thursday" in exp_lower:
        result = base - timedelta(days=base.weekday() + 4)
    elif "last tuesday" in exp_lower:
        result = base - timedelta(days=base.weekday() + 6)
    elif "next tuesday" in exp_lower:
        result = base + timedelta(days=(7 - base.weekday()) + 1)
    elif "next thursday" in exp_lower:
        result = base + timedelta(days=(7 - base.weekday()) + 3)
    elif "去年今天" in expression or "去年今日" in expression or "last year" in exp_lower:
        try:
            result = base.replace(year=base.year - 1)
        except ValueError:
            result = base.replace(year=base.year - 1, day=28)
    elif "明年今天" in expression or "明年今日" in expression or "next year" in exp_lower:
        try:
            result = base.replace(year=base.year + 1)
        except ValueError:
            result = base.replace(year=base.year + 1, day=28)
    elif backward_workday_offset:
        result = _add_workdays(base, -_chinese_number_to_int(backward_workday_offset.group(1)))
    elif workday_offset:
        result = _add_workdays(base, _chinese_number_to_int(workday_offset.group(1)))
    elif hour_offset:
        result = base + timedelta(hours=_chinese_number_to_int(hour_offset.group(1)))
    elif "半个月前" in expression or "半月前" in expression:
        result = base - timedelta(days=15)
    elif "半个月后" in expression or "半月后" in expression:
        result = base + timedelta(days=15)
    elif backward_days:
        result = base - timedelta(days=_chinese_number_to_int(backward_days.group(1)))
    elif forward_days:
        result = base + timedelta(days=_chinese_number_to_int(forward_days.group(1)))
    elif natural_days:
        result = base + timedelta(days=_chinese_number_to_int(natural_days.group(1)))
    elif "两周后" in expression:
        result = base + timedelta(weeks=2)
    elif backward_weeks:
        result = base - timedelta(weeks=_chinese_number_to_int(backward_weeks.group(1)))
    elif forward_weeks:
        result = base + timedelta(weeks=_chinese_number_to_int(forward_weeks.group(1)))
    elif backward_month_offset:
        result = _add_months(base, -_chinese_number_to_int(backward_month_offset.group(1)))
    elif month_offset:
        result = _add_months(base, _chinese_number_to_int(month_offset.group(1)))
    elif re.search(rf"(?<!哪){_DATE_NUMBER_PATTERN}\s*天(?![是为])", expression):
        day_match = re.search(rf"(?<!哪){_DATE_NUMBER_PATTERN}\s*天(?![是为])", expression)
        result = base + timedelta(days=_chinese_number_to_int(day_match.group(1)))
    elif re.search(rf"{_DATE_NUMBER_PATTERN}\s*年后", expression):
        years = _chinese_number_to_int(re.search(rf"{_DATE_NUMBER_PATTERN}\s*年后", expression).group(1))
        result = base.replace(year=base.year + years)
    elif "儿童节" in expression:
        result = base.replace(month=6, day=1, hour=0, minute=0)
    elif "圣诞节" in expression:
        result = base.replace(month=12, day=25, hour=0, minute=0)
    elif "元旦" in expression:
        result = base.replace(month=1, day=1, hour=0, minute=0)
    elif month_day_target:
        result = base.replace(
            month=int(month_day_target.group(1)),
            day=int(month_day_target.group(2)),
            hour=0,
            minute=0,
        )
    elif explicit_dates:
        match = explicit_dates[-1]
        result = _parse_date_datetime(
            expression,
            int(match.group(1)),
            int(match.group(2)),
            int(match.group(3)),
            match.end(),
        )
        trailing_month_days = list(re.finditer(r"(\d{1,2})月(\d{1,2})日", expression))
        if trailing_month_days and trailing_month_days[-1].start() > match.end():
            target = trailing_month_days[-1]
            result = result.replace(
                month=int(target.group(1)),
                day=int(target.group(2)),
                hour=0,
                minute=0,
            )
    else:
        european = re.search(r"(?<!\d)(\d{1,2})/(\d{1,2})/(\d{4})(?!\d)", expression)
        result = (
            datetime(
                int(european.group(3)),
                int(european.group(2)),
                int(european.group(1)),
            )
            if european
            else base
        )

    if result:
        return _date_result_payload(expression, base, result)
    return {"error": "Could not parse date expression", "expression": expression}


def _read_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown", ".csv", ".json", ".log"}:
        return path.read_text(encoding="utf-8-sig", errors="replace")
    if suffix == ".docx":
        return _read_docx_text(path)
    return path.read_text(encoding="utf-8-sig", errors="replace")


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document

        doc = Document(path)
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return _read_docx_zip_text(path)


def _read_docx_zip_text(path: Path) -> str:
    import xml.etree.ElementTree as ET

    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraphs: list[str] = []
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        candidates = ["word/document.xml", *sorted(name for name in names if name.endswith(".xml"))]
        seen: set[str] = set()
        for candidate in candidates:
            if candidate in seen or candidate not in names:
                continue
            seen.add(candidate)
            try:
                root = ET.parse(zf.open(candidate)).getroot()
            except ET.ParseError:
                continue
            texts = [
                elem.text.strip()
                for elem in root.iter(f"{{{ns}}}t")
                if elem.text and elem.text.strip()
            ]
            if texts:
                paragraphs.append("".join(texts))
                if candidate == "word/document.xml":
                    break
    return "\n".join(paragraphs)


def _document_chunks(text: str, *, max_chars: int = 1200, overlap: int = 160) -> list[tuple[str, str]]:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    chunks: list[tuple[str, str]] = []
    current_title = "Document"
    current: list[str] = []

    def flush() -> None:
        content = "\n".join(part for part in current if part).strip()
        if not content:
            return
        if len(content) <= max_chars:
            chunks.append((current_title, content))
            return
        start = 0
        while start < len(content):
            piece = content[start : start + max_chars].strip()
            if piece:
                chunks.append((current_title, piece))
            if start + max_chars >= len(content):
                break
            start += max(1, max_chars - overlap)

    for line in lines:
        if not line:
            continue
        is_heading = (
            line.startswith("#")
            or bool(re.match(r"^(?:第[一二三四五六七八九十百\d]+[章节条]|[一二三四五六七八九十百\d]+[、.．])", line))
        )
        if is_heading:
            flush()
            current_title = line.lstrip("#").strip()[:120] or "Document"
            current = []
        else:
            current.append(line)
    flush()
    if not chunks and text.strip():
        chunks.append(("Document", text.strip()[:max_chars]))
    return chunks


def _document_query_terms(query: str) -> list[str]:
    terms: list[str] = []
    for term in re.findall(r"[A-Za-z0-9_./-]{2,}", query.lower()):
        terms.append(term)
    for seq in re.findall(r"[\u4e00-\u9fff]{2,}", query):
        terms.append(seq)
        max_n = min(6, len(seq))
        for size in range(2, max_n + 1):
            for start in range(0, len(seq) - size + 1):
                terms.append(seq[start : start + size])
    deduped: list[str] = []
    seen: set[str] = set()
    for term in terms:
        if term not in seen:
            seen.add(term)
            deduped.append(term)
    return deduped[:240]


def _score_document_chunk(query: str, terms: list[str], title: str, content: str) -> int:
    haystack = f"{title}\n{content}"
    lower_haystack = haystack.lower()
    lower_query = query.lower().strip()
    score = 0
    if lower_query and lower_query in lower_haystack:
        score += 1000
    for term in terms:
        count = lower_haystack.count(term.lower())
        if not count:
            continue
        if re.search(r"[\u4e00-\u9fff]", term):
            score += count * max(2, len(term))
        else:
            score += count * 6
    return score


def _highlight_snippet(content: str, terms: list[str], max_chars: int) -> str:
    if len(content) <= max_chars:
        return content
    best_pos = 0
    for term in sorted(terms, key=len, reverse=True):
        if len(term) < 2:
            continue
        pos = content.lower().find(term.lower())
        if pos >= 0:
            best_pos = max(0, pos - max_chars // 3)
            break
    snippet = content[best_pos : best_pos + max_chars].strip()
    if best_pos > 0:
        snippet = "..." + snippet
    if best_pos + max_chars < len(content):
        snippet += "..."
    return snippet


def register_tools(*, register_tool: Callable[..., Callable], object_schema: Callable[..., dict[str, Any]]) -> None:
    """Register contestant MCP-style tools.

    This file is loaded by source/toolkits/main_mcp.py. Contestants can add,
    remove, or replace tools here when a capability is better exposed as a
    direct MCP-style function than as a SKILL.md package.
    """

    @register_tool(
        name="mock_order_lookup",
        description="Mock MCP-style tool. Returns a fixed mock order lookup result for demo purposes.",
        input_schema=object_schema(
            {
                "order_id": {
                    "type": "string",
                    "description": "Mock order id, for example MOCK-1001.",
                }
            },
            ["order_id"],
        ),
        kind="mcp",
        risk="low",
    )
    def mock_order_lookup(order_id: str) -> str:
        return json.dumps(
            {
                "mock_result": "mock-order-lookup-ok",
                "source": "mock_mcp",
                "order_id": order_id,
            },
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="mock_policy_check",
        description="Mock MCP-style tool. Returns a fixed mock policy check result for demo purposes.",
        input_schema=object_schema(
            {
                "payload": {
                    "type": "string",
                    "description": "Mock payload to check.",
                }
            },
            ["payload"],
        ),
        kind="mcp",
        risk="low",
    )
    def mock_policy_check(payload: str) -> str:
        return json.dumps(
            {
                "mock_result": "mock-policy-check-ok",
                "source": "mock_mcp",
                "payload_preview": payload[:80],
            },
            ensure_ascii=False,
            indent=2,
        )

    # =========================================================================
    # P0 Competition Tools
    # =========================================================================

    @register_tool(
        name="evidence_chain_analyze",
        description=(
            "Parse frontend/backend logs, HAR files, screenshots, and a validation "
            "schema; correlate complete foreground failure flows and map root causes."
        ),
        input_schema=object_schema(
            {
                "frontend_paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "default": [],
                },
                "backend_paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "default": [],
                },
                "har_paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "default": [],
                },
                "schema_path": {"type": "string"},
                "screenshot_paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "default": [],
                },
            },
            ["schema_path"],
        ),
        kind="mcp",
        risk="medium",
    )
    def evidence_chain_analyze(
        frontend_paths: list[str] | None = None,
        backend_paths: list[str] | None = None,
        har_paths: list[str] | None = None,
        schema_path: str = "",
        screenshot_paths: list[str] | None = None,
    ) -> str:
        return json.dumps(
            analyze_evidence_chain(
                frontend_paths=frontend_paths,
                backend_paths=backend_paths,
                har_paths=har_paths,
                schema_path=schema_path,
                screenshot_paths=screenshot_paths,
            ),
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="api_test_execute",
        description=(
            "Validate and sequentially execute a complete HTTP API test plan. "
            "Supports response variable extraction and deterministic JSON assertions."
        ),
        input_schema=object_schema(
            {
                "plan": {
                    "type": "object",
                    "description": (
                        "Execution plan with base_url and ordered cases/steps. "
                        "Each case may contain assert/expectedStatus/expectedFields/"
                        "expectedValues; steps may extract variables by JSON path."
                    ),
                    "additionalProperties": True,
                }
            },
            ["plan"],
        ),
        kind="mcp",
        risk="high",
    )
    def api_test_execute(
        plan: dict[str, Any],
        package_id: str = "",
        auth_config: dict[str, Any] | None = None,
    ) -> str:
        return json.dumps(
            execute_api_test_plan(
                plan=plan,
                package_id=package_id,
                auth_config=auth_config,
            ),
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="document_search",
        description=(
            "Search relevant snippets across declared DOCX/Markdown/text documents. "
            "Use for coding-standard Q&A over multiple attached specification files."
        ),
        input_schema=object_schema(
            {
                "paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Declared document paths or directories containing .docx/.md/.txt files.",
                },
                "query": {
                    "type": "string",
                    "description": "Question or keywords to search for.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Maximum snippets to return.",
                    "default": 8,
                },
                "max_chars": {
                    "type": "integer",
                    "description": "Maximum characters per returned snippet.",
                    "default": 1200,
                },
            },
            ["paths", "query"],
        ),
        kind="mcp",
        risk="low",
    )
    def document_search(
        paths: list[str],
        query: str,
        limit: int = 8,
        max_chars: int = 1200,
    ) -> str:
        terms = _document_query_terms(query)
        scored: list[dict[str, Any]] = []
        errors: list[str] = []
        supported = {".docx", ".md", ".markdown", ".txt", ".log"}

        for raw_path in paths or []:
            path = Path(raw_path)
            candidates = [path]
            if path.is_dir():
                candidates = [
                    child
                    for child in path.rglob("*")
                    if child.is_file() and child.suffix.lower() in supported
                ]

            for candidate in candidates:
                if candidate.suffix.lower() not in supported:
                    continue
                try:
                    text = _read_document_text(candidate)
                except Exception as exc:
                    errors.append(f"{candidate}: {type(exc).__name__}: {exc}")
                    continue
                for section_index, (title, content) in enumerate(_document_chunks(text)):
                    score = _score_document_chunk(query, terms, title, content)
                    if score <= 0:
                        continue
                    scored.append(
                        {
                            "score": score,
                            "path": str(candidate),
                            "filename": candidate.name,
                            "section": title,
                            "section_index": section_index,
                            "snippet": _highlight_snippet(content, terms, max(200, min(max_chars, 4000))),
                        }
                    )

        scored.sort(key=lambda item: item["score"], reverse=True)
        selected = scored[: max(1, min(limit, 20))]
        return json.dumps(
            {
                "query": query,
                "match_count": len(scored),
                "results": selected,
                "errors": errors[:10],
            },
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="http_request",
        description="Send HTTP request and return response. Supports GET/POST/PUT/DELETE with custom headers and body.",
        input_schema=object_schema(
            {
                "url": {
                    "type": "string",
                    "description": "Target URL (e.g., http://localhost:18080/api/question)",
                },
                "method": {
                    "type": "string",
                    "description": "HTTP method: GET, POST, PUT, DELETE, PATCH",
                    "default": "GET",
                },
                "headers": {
                    "type": "object",
                    "description": "Request headers as key-value pairs",
                    "default": {},
                },
                "body": {
                    "type": "string",
                    "description": "Request body (for POST/PUT/PATCH)",
                    "default": "",
                },
                "timeout": {
                    "type": "integer",
                    "description": "Timeout in seconds",
                    "default": 30,
                },
            },
            ["url"],
        ),
        kind="mcp",
        risk="medium",
    )
    def http_request(
        url: str,
        method: str = "GET",
        headers: dict[str, str] | None = None,
        body: str = "",
        timeout: int = 30,
    ) -> str:
        """Send HTTP request and return response status + body."""
        headers = headers or {}
        method = method.upper()

        try:
            req = urllib.request.Request(url, method=method)
            for key, value in headers.items():
                req.add_header(key, value)

            if body and method in ("POST", "PUT", "PATCH"):
                if isinstance(body, str):
                    body = body.encode("utf-8")
                req.data = body

            with urllib.request.urlopen(req, timeout=timeout) as response:
                status = response.status
                response_headers = dict(response.headers)
                response_body = response.read().decode("utf-8", errors="replace")

            return json.dumps(
                {
                    "status": status,
                    "headers": response_headers,
                    "body": response_body[:50000],
                    "success": 200 <= status < 300,
                },
                ensure_ascii=False,
                indent=2,
            )
        except urllib.error.HTTPError as exc:
            return json.dumps(
                {
                    "status": exc.code,
                    "error": str(exc.reason),
                    "body": exc.read().decode("utf-8", errors="replace")[:10000],
                    "success": False,
                },
                ensure_ascii=False,
                indent=2,
            )
        except urllib.error.URLError as exc:
            return json.dumps(
                {
                    "status": 0,
                    "error": str(exc.reason),
                    "success": False,
                },
                ensure_ascii=False,
                indent=2,
            )
        except Exception as exc:
            return json.dumps(
                {
                    "status": -1,
                    "error": str(exc),
                    "success": False,
                },
                ensure_ascii=False,
                indent=2,
            )

    @register_tool(
        name="zip_extract",
        description="Extract ZIP file contents. Returns list of extracted files.",
        input_schema=object_schema(
            {
                "zip_path": {
                    "type": "string",
                    "description": "Path to ZIP file",
                },
                "output_dir": {
                    "type": "string",
                    "description": "Output directory (default: temp directory)",
                    "default": "",
                },
            },
            ["zip_path"],
        ),
        kind="mcp",
        risk="medium",
    )
    def zip_extract(zip_path: str, output_dir: str = "") -> str:
        """Extract ZIP file and return list of extracted files."""
        zip_path = Path(zip_path)
        if not zip_path.exists():
            return json.dumps({"error": f"ZIP file not found: {zip_path}"}, ensure_ascii=False)

        if not output_dir:
            output_dir = tempfile.mkdtemp(prefix="zip_extract_")
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)

        extracted = []
        try:
            with zipfile.ZipFile(zip_path, "r") as zf:
                output_root = Path(output_dir).resolve()
                for member in zf.infolist():
                    target = _safe_archive_target(output_root, member.filename)
                    mode = member.external_attr >> 16
                    if stat.S_ISLNK(mode):
                        raise ValueError(f"Unsafe archive member (link): {member.filename}")
                    if member.is_dir():
                        target.mkdir(parents=True, exist_ok=True)
                        continue
                    target.parent.mkdir(parents=True, exist_ok=True)
                    with zf.open(member, "r") as source, target.open("wb") as destination:
                        shutil.copyfileobj(source, destination)
                    extracted.append(str(target))
        except zipfile.BadZipFile:
            return json.dumps({"error": f"Invalid ZIP file: {zip_path}"}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"error": str(exc)}, ensure_ascii=False)

        return json.dumps(
            {
                "output_dir": str(output_dir),
                "files": extracted,
                "count": len(extracted),
            },
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="tar_extract",
        description="Extract TAR/TAR.GZ/TAR.BZ2 file contents. Returns list of extracted files. Supports .tar, .tar.gz, .tgz, .tar.bz2 formats.",
        input_schema=object_schema(
            {
                "tar_path": {
                    "type": "string",
                    "description": "Path to TAR file (.tar, .tar.gz, .tgz, .tar.bz2)",
                },
                "output_dir": {
                    "type": "string",
                    "description": "Output directory (default: temp directory)",
                    "default": "",
                },
            },
            ["tar_path"],
        ),
        kind="mcp",
        risk="medium",
    )
    def tar_extract(tar_path: str, output_dir: str = "") -> str:
        """Extract TAR file and return list of extracted files."""
        tar_path = Path(tar_path)
        if not tar_path.exists():
            return json.dumps({"error": f"TAR file not found: {tar_path}"}, ensure_ascii=False)

        if not output_dir:
            output_dir = tempfile.mkdtemp(prefix="tar_extract_")
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)

        extracted = []
        try:
            with tarfile.open(tar_path, "r:*") as tf:
                output_root = Path(output_dir).resolve()
                for member in tf.getmembers():
                    target = _safe_archive_target(output_root, member.name)
                    if member.issym() or member.islnk() or member.isdev():
                        raise ValueError(f"Unsafe archive member (link/device): {member.name}")
                    if member.isdir():
                        target.mkdir(parents=True, exist_ok=True)
                        continue
                    if not member.isfile():
                        raise ValueError(f"Unsafe archive member type: {member.name}")
                    source = tf.extractfile(member)
                    if source is None:
                        raise ValueError(f"Could not read archive member: {member.name}")
                    target.parent.mkdir(parents=True, exist_ok=True)
                    with source, target.open("wb") as destination:
                        shutil.copyfileobj(source, destination)
                    extracted.append(str(target))
        except tarfile.TarError as exc:
            return json.dumps({"error": f"Invalid TAR file: {exc}"}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"error": str(exc)}, ensure_ascii=False)

        return json.dumps(
            {
                "output_dir": str(output_dir),
                "files": extracted,
                "count": len(extracted),
            },
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="csv_read",
        description="Read one or more CSV files and return JSON arrays. Supports custom delimiter and encoding.",
        input_schema=object_schema(
            {
                "path": {
                    "type": "string",
                    "description": "Path to CSV file for single-file mode",
                },
                "delimiter": {
                    "type": "string",
                    "description": "Column delimiter",
                    "default": ",",
                },
                "encoding": {
                    "type": "string",
                    "description": "File encoding",
                    "default": "utf-8",
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum rows to read (0 = all)",
                    "default": 0,
                },
                "items": {
                    "type": "array",
                    "description": "Optional batch items. Each item may contain id, path, delimiter, encoding, and max_rows.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "path": {"type": "string"},
                            "delimiter": {"type": "string"},
                            "encoding": {"type": "string"},
                            "max_rows": {"type": "integer"},
                        },
                    },
                    "default": [],
                },
            },
            [],
        ),
        kind="mcp",
        risk="low",
    )
    def csv_read(
        path: str = "",
        delimiter: str = ",",
        encoding: str = "utf-8",
        max_rows: int = 0,
        items: list[dict[str, Any]] | None = None,
    ) -> str:
        """Read one or more CSV files and return JSON arrays."""

        def read_one(
            item_path: str,
            item_delimiter: str = ",",
            item_encoding: str = "utf-8",
            item_max_rows: int = 0,
        ) -> dict[str, Any]:
            csv_path = Path(item_path)
            if not csv_path.exists():
                return {"error": f"File not found: {csv_path}", "path": str(csv_path)}

            rows = []
            try:
                with csv_path.open("r", encoding=item_encoding, newline="") as f:
                    reader = csv.DictReader(f, delimiter=item_delimiter)
                    for i, row in enumerate(reader):
                        if item_max_rows > 0 and i >= item_max_rows:
                            break
                        rows.append(row)
            except Exception as exc:
                return {"error": str(exc), "path": str(csv_path)}

            return {
                "path": str(csv_path),
                "rows": len(rows),
                "columns": list(rows[0].keys()) if rows else [],
                "data": rows[:1000],
            }

        if items:
            results = []
            for index, item in enumerate(items):
                payload = read_one(
                    str(item.get("path", "")),
                    str(item.get("delimiter") or delimiter),
                    str(item.get("encoding") or encoding),
                    int(item.get("max_rows", max_rows) or 0),
                )
                if "id" in item:
                    payload["id"] = item["id"]
                else:
                    payload["index"] = index
                results.append(payload)
            return json.dumps({"results": results, "count": len(results)}, ensure_ascii=False, indent=2)

        if not path:
            return json.dumps({"error": "path or items is required"}, ensure_ascii=False)
        return json.dumps(
            read_one(path, delimiter, encoding, max_rows),
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="csv_aggregate",
        description="Aggregate CSV data: SUM/AVG/COUNT/MIN/MAX with optional GROUP BY.",
        input_schema=object_schema(
            {
                "data": {
                    "type": "array",
                    "description": "CSV rows as JSON array (from csv_read)",
                },
                "operation": {
                    "type": "string",
                    "description": "Aggregation: SUM, AVG, COUNT, MIN, MAX",
                },
                "column": {
                    "type": "string",
                    "description": "Target column name",
                },
                "group_by": {
                    "type": "string",
                    "description": "Group by column (optional)",
                    "default": "",
                },
            },
            ["data", "operation", "column"],
        ),
        kind="mcp",
        risk="low",
    )
    def csv_aggregate(
        data: list[dict[str, Any]],
        operation: str,
        column: str,
        group_by: str = "",
    ) -> str:
        """Aggregate CSV data."""
        if not data:
            return json.dumps({"error": "Empty data"}, ensure_ascii=False)

        operation = operation.upper()
        values = []
        for row in data:
            try:
                val = float(row.get(column, 0))
                values.append(val)
            except (ValueError, TypeError):
                pass

        if not values:
            return json.dumps({"error": f"No numeric values in column: {column}"}, ensure_ascii=False)

        if operation == "SUM":
            result = sum(values)
        elif operation == "AVG":
            result = sum(values) / len(values)
        elif operation == "COUNT":
            result = len(values)
        elif operation == "MIN":
            result = min(values)
        elif operation == "MAX":
            result = max(values)
        else:
            return json.dumps({"error": f"Unknown operation: {operation}"}, ensure_ascii=False)

        return json.dumps(
            {
                "operation": operation,
                "column": column,
                "result": result,
                "count": len(values),
            },
            ensure_ascii=False,
            indent=2,
        )

    @register_tool(
        name="code_execute",
        description="Execute code (Python/Java/Node.js) and return stdout/stderr/exit_code.",
        input_schema=object_schema(
            {
                "language": {
                    "type": "string",
                    "description": "Programming language: python, java, node",
                },
                "code": {
                    "type": "string",
                    "description": "Source code to execute",
                },
                "stdin": {
                    "type": "string",
                    "description": "Standard input",
                    "default": "",
                },
                "args": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Command-line arguments passed to the program",
                    "default": [],
                },
                "stdin_cases": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional multiple stdin cases. Java is compiled once and run once per case.",
                    "default": [],
                },
                "timeout": {
                    "type": "integer",
                    "description": "Execution timeout in seconds",
                    "default": 30,
                },
            },
            ["language", "code"],
        ),
        kind="mcp",
        risk="high",
    )
    def code_execute(
        language: str,
        code: str,
        stdin: str = "",
        args: list[str] | None = None,
        stdin_cases: list[str] | None = None,
        timeout: int = 30,
    ) -> str:
        """Execute code and return output."""
        language = language.lower()
        args = [str(value) for value in (args or [])]
        stdin_cases = [str(value) for value in (stdin_cases or [])]

        if language == "python":
            cmd = [sys.executable, "-c", code, *args]
        elif language == "java":
            class_name = _java_main_class_name(code)
            if not class_name:
                return json.dumps(
                    {
                        "language": language,
                        "exit_code": -1,
                        "stdout": "",
                        "stderr": "Could not identify a Java class containing main().",
                        "error": "main class not found",
                    },
                    ensure_ascii=False,
                    indent=2,
                )
            try:
                with tempfile.TemporaryDirectory(prefix="java_execute_") as temp_dir:
                    java_file = Path(temp_dir) / f"{class_name}.java"
                    java_file.write_text(code, encoding="utf-8")
                    compile_result = subprocess.run(
                        ["javac", "-encoding", "UTF-8", java_file.name],
                        cwd=temp_dir,
                        capture_output=True,
                        text=True,
                        encoding="utf-8",
                        errors="replace",
                        timeout=timeout,
                    )
                    if compile_result.returncode != 0:
                        return json.dumps(
                            {
                                "language": language,
                                "exit_code": -1,
                                "stdout": "",
                                "stderr": compile_result.stderr,
                                "error": "Compilation failed",
                            },
                            ensure_ascii=False,
                            indent=2,
                        )
                    cases = stdin_cases if stdin_cases else [stdin]
                    runs = []
                    java_cmd = [
                        "java",
                        "-Dfile.encoding=UTF-8",
                        "-Dsun.stdout.encoding=UTF-8",
                        "-Dsun.stderr.encoding=UTF-8",
                        "-cp",
                        temp_dir,
                        class_name,
                        *args,
                    ]
                    for index, case_stdin in enumerate(cases):
                        run_result = subprocess.run(
                            java_cmd,
                            input=case_stdin,
                            capture_output=True,
                            text=True,
                            encoding="utf-8",
                            errors="replace",
                            timeout=timeout,
                        )
                        runs.append(
                            {
                                "index": index,
                                "stdin": case_stdin,
                                "args": args,
                                "exit_code": run_result.returncode,
                                "stdout": run_result.stdout[:50000],
                                "stderr": run_result.stderr[:10000],
                            }
                        )
                    if stdin_cases:
                        return json.dumps(
                            {
                                "language": language,
                                "class_name": class_name,
                                "compile_exit_code": compile_result.returncode,
                                "runs": runs,
                            },
                            ensure_ascii=False,
                            indent=2,
                        )
                    run_result = runs[0]
                    return json.dumps(
                        {
                            "language": language,
                            "class_name": class_name,
                            "exit_code": run_result["exit_code"],
                            "stdout": run_result["stdout"],
                            "stderr": run_result["stderr"],
                        },
                        ensure_ascii=False,
                        indent=2,
                    )
            except subprocess.TimeoutExpired:
                return json.dumps(
                    {
                        "language": language,
                        "exit_code": -1,
                        "stdout": "",
                        "stderr": f"Execution timeout after {timeout}s",
                        "error": "timeout",
                    },
                    ensure_ascii=False,
                    indent=2,
                )
            except Exception as exc:
                return json.dumps(
                    {
                        "language": language,
                        "exit_code": -1,
                        "stdout": "",
                        "stderr": str(exc),
                        "error": str(exc),
                    },
                    ensure_ascii=False,
                    indent=2,
                )
        elif language == "node":
            cmd = ["node", "-e", code, *args]
        else:
            return json.dumps({"error": f"Unsupported language: {language}"}, ensure_ascii=False)

        try:
            result = subprocess.run(
                cmd,
                input=stdin,
                capture_output=True,
                text=True,
                timeout=timeout,
            )
            return json.dumps(
                {
                    "language": language,
                    "exit_code": result.returncode,
                    "stdout": result.stdout[:50000],
                    "stderr": result.stderr[:10000],
                },
                ensure_ascii=False,
                indent=2,
            )
        except subprocess.TimeoutExpired:
            return json.dumps(
                {
                    "language": language,
                    "exit_code": -1,
                    "stdout": "",
                    "stderr": f"Execution timeout after {timeout}s",
                    "error": "timeout",
                },
                ensure_ascii=False,
                indent=2,
            )
        except Exception as exc:
            return json.dumps(
                {
                    "language": language,
                    "exit_code": -1,
                    "stdout": "",
                    "stderr": str(exc),
                    "error": str(exc),
                },
                ensure_ascii=False,
                indent=2,
            )

    # =========================================================================
    # P0 Competition Tools: date_compute, sql_query
    # =========================================================================

    @register_tool(
        name="date_compute",
        description="Parse natural-language date sentence(s) and compute target dates. For multiple date questions, pass items in one call to avoid repeated LLM/tool round trips.",
        input_schema=object_schema(
            {
                "expression": {
                    "type": "string",
                    "description": "Complete original sentence for single-item mode, e.g. '今天是2026年5月6日，明天是几号'. Do not shorten away hours or week anchors.",
                },
                "base_date": {
                    "type": "string",
                    "description": "Base date in YYYY-MM-DD format for relative date calculations",
                    "default": "",
                },
                "items": {
                    "type": "array",
                    "description": "Optional batch items. Each item may contain id, expression, and base_date. Use this for multiple independent date questions.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "expression": {"type": "string"},
                            "base_date": {"type": "string"},
                        },
                    },
                    "default": [],
                },
            },
            [],
        ),
        kind="mcp",
        risk="low",
    )
    def date_compute(
        expression: str = "",
        base_date: str = "",
        items: list[dict[str, Any]] | None = None,
    ) -> str:
        """Parse one or many natural language date expressions."""
        try:
            if items:
                results = []
                for index, item in enumerate(items):
                    item_expression = str(item.get("expression", "")).strip()
                    item_payload = _compute_date_payload(
                        item_expression,
                        str(item.get("base_date") or base_date or ""),
                    )
                    if "id" in item:
                        item_payload["id"] = item["id"]
                    else:
                        item_payload["index"] = index
                    results.append(item_payload)
                return json.dumps(
                    {"results": results, "count": len(results)},
                    ensure_ascii=False,
                    indent=2,
                )

            if not expression:
                return json.dumps(
                    {"error": "expression or items is required"},
                    ensure_ascii=False,
                )
            return json.dumps(
                _compute_date_payload(expression, base_date),
                ensure_ascii=False,
                indent=2,
            )
        except Exception as exc:
            return json.dumps({"error": str(exc)}, ensure_ascii=False)

    @register_tool(
        name="sql_query",
        description="Execute one or more SELECT queries on SQLite database(s) and return results as JSON. Only SELECT queries allowed.",
        input_schema=object_schema(
            {
                "db_path": {
                    "type": "string",
                    "description": "Path to SQLite database file for single-query mode or shared batch default",
                },
                "query": {
                    "type": "string",
                    "description": "SQL query to execute in single-query mode (SELECT only for safety)",
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum rows to return",
                    "default": 100,
                },
                "items": {
                    "type": "array",
                    "description": "Optional batch items. Each item may contain id, db_path, query, and max_rows.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "db_path": {"type": "string"},
                            "query": {"type": "string"},
                            "max_rows": {"type": "integer"},
                        },
                    },
                    "default": [],
                },
            },
            [],
        ),
        kind="mcp",
        risk="medium",
    )
    def sql_query(
        db_path: str = "",
        query: str = "",
        max_rows: int = 100,
        items: list[dict[str, Any]] | None = None,
    ) -> str:
        """Execute one or more SELECT queries on SQLite database(s)."""

        def run_one(item_db_path: str, item_query: str, item_max_rows: int = 100) -> dict[str, Any]:
            try:
                import sqlite3
                path = Path(item_db_path)
                if not path.exists():
                    return {"error": f"Database not found: {item_db_path}", "db_path": str(path)}

                query_stripped = item_query.strip().upper()
                if not query_stripped.startswith("SELECT"):
                    return {"error": "Only SELECT queries are allowed", "sql": item_query}

                conn = sqlite3.connect(str(path))
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                cursor.execute(item_query)
                rows = cursor.fetchmany(item_max_rows)
                columns = [desc[0] for desc in cursor.description] if cursor.description else []

                result_data = []
                for row in rows:
                    result_data.append({col: row[i] for i, col in enumerate(columns)})

                conn.close()
                return {
                    "columns": columns,
                    "rows": result_data,
                    "count": len(result_data),
                    "sql": item_query,
                    "db_path": str(path),
                }
            except Exception as exc:
                return {"error": str(exc), "sql": item_query, "db_path": item_db_path}

        if items:
            results = []
            for index, item in enumerate(items):
                payload = run_one(
                    str(item.get("db_path") or db_path),
                    str(item.get("query", "")),
                    int(item.get("max_rows", max_rows) or max_rows),
                )
                if "id" in item:
                    payload["id"] = item["id"]
                else:
                    payload["index"] = index
                results.append(payload)
            return json.dumps({"results": results, "count": len(results)}, ensure_ascii=False, indent=2, default=str)

        if not db_path or not query:
            return json.dumps({"error": "db_path and query, or items, are required"}, ensure_ascii=False)
        return json.dumps(run_one(db_path, query, max_rows), ensure_ascii=False, indent=2, default=str)

    @register_tool(
        name="image_read",
        description="Read one or more image files and return base64 for multimodal analysis. Supports PNG, JPG, JPEG, BMP, GIF, WebP. For multiple images, pass items in one call.",
        input_schema=object_schema(
            {
                "path": {
                    "type": "string",
                    "description": "Path to one image file for single-image mode",
                },
                "question": {
                    "type": "string",
                    "description": "Question about the image to ask the vision model (e.g., 'What text is shown?', 'Describe the error message').",
                    "default": "",
                },
                "items": {
                    "type": "array",
                    "description": "Optional batch image items. Each item may contain id, path, and question.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "path": {"type": "string"},
                            "question": {"type": "string"},
                        },
                    },
                    "default": [],
                },
            },
            [],
        ),
        kind="mcp",
        risk="low",
    )
    def image_read(
        path: str = "",
        question: str = "",
        items: list[dict[str, Any]] | None = None,
    ) -> str:
        """Read one or more image files and return base64 data for multimodal LLM calls."""
        import base64

        mime_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".bmp": "image/bmp",
            ".gif": "image/gif",
            ".webp": "image/webp",
        }

        def read_one(
            img_path_text: str,
            item_question: str = "",
            item_id: Any = None,
            item_index: int | None = None,
        ) -> dict[str, Any]:
            img_path = Path(img_path_text)
            if not img_path.exists():
                return {"error": f"Image file not found: {img_path_text}", "path": img_path_text}

            suffix = img_path.suffix.lower()
            mime_type = mime_map.get(suffix, "image/png")
            if suffix not in mime_map:
                return {"error": f"Unsupported image format: {suffix}", "path": str(img_path)}

            with open(img_path, "rb") as f:
                img_bytes = f.read()

            if len(img_bytes) > 10 * 1024 * 1024:
                return {"error": "Image too large (>10MB)", "path": str(img_path)}

            label_parts = []
            if item_index is not None:
                label_parts.append(f"第 {item_index} 张")
            label_parts.append(f"文件名 {img_path.name}")
            default_question = (
                "请详细描述这张图片的内容，包括所有可见的文字、数字、表格、图表、错误信息等；"
                f"保持与{'，'.join(label_parts)}的对应关系。"
            )
            payload: dict[str, Any] = {
                "__image__": True,
                "path": str(img_path),
                "mime_type": mime_type,
                "base64": base64.b64encode(img_bytes).decode("utf-8"),
                "question": item_question or default_question,
            }
            if item_id is not None:
                payload["id"] = item_id
            return payload

        try:
            if items:
                images = [
                    read_one(
                        str(item.get("path", "")),
                        str(item.get("question") or question or ""),
                        item.get("id"),
                        index,
                    )
                    for index, item in enumerate(items, start=1)
                ]
                return json.dumps(
                    {
                        "__images__": True,
                        "images": images,
                        "count": len(images),
                    },
                    ensure_ascii=False,
                )

            if not path:
                return json.dumps({"error": "path or items is required"}, ensure_ascii=False)
            return json.dumps(read_one(path, question, item_index=1), ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"error": f"Image read failed: {exc}"}, ensure_ascii=False)
