"""DOCX document parser using python-docx.

Adapted from Document-Logic-MCP with section_ref support for FTS5 indexing.
"""

import logging
import re
from pathlib import Path
from typing import List, Optional
from docx import Document
from .base import BaseParser, ParseResult, Section
from .docx_table_converter import get_table_paragraph_positions, table_to_markdown
from .pdf_parser import _is_heading, _make_section_ref, _make_parent_ref

logger = logging.getLogger(__name__)

# Pattern-based heading detection (fallback when no Heading styles exist).
_NUMBERED_HEADING_PATTERNS = [
    (re.compile(r'^\d+\.\d+\.\d+\s+[A-Z]\w'), 3),
    (re.compile(r'^\d+\.\d+\s+[A-Z]\w'), 2),
    (re.compile(r'^\d+\.\s+[A-Z]\w'), 1),
]

_MAX_HEADING_LENGTH = 120
_LIST_ITEM_CHARS = re.compile(r'[,;]')

_SENTENCE_STARTERS = frozenset({
    'The', 'This', 'These', 'Those', 'That',
    'A', 'An', 'Each', 'Every', 'All',
    'It', 'Its', 'Our', 'We', 'They',
    'When', 'If', 'For', 'In', 'On', 'At',
    'There', 'Here', 'Any', 'Some', 'No',
    'Most', 'Many', 'Several',
})

_NUMBER_PREFIX_RE = re.compile(r'^\d+(?:\.\d+)*\.?\s+(.*)')

# XML namespace for Word documents
_WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006}'

# Formatting-based heading detection constants
_MAX_FMT_HEADING_LEN = 80
_MAX_ALLCAPS_LEN = 60
_MAX_ALLCAPS_WORDS = 8
_MIN_FONT_DELTA_PT = 2.0
_BOLD_MAJORITY_THRESHOLD = 0.5


def _compute_body_font_size(doc) -> Optional[float]:
    """Compute the weighted median font size of body text paragraphs."""
    samples: list[tuple[float, int]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            continue
        for run in para.runs:
            if run.font.size is not None and len(run.text.strip()) > 0:
                size_pt = run.font.size.pt
                char_count = len(run.text)
                samples.append((size_pt, char_count))

    if not samples:
        return None

    total_chars = sum(c for _, c in samples)
    samples.sort(key=lambda s: s[0])
    cumulative = 0
    half = total_chars / 2.0
    for size_pt, char_count in samples:
        cumulative += char_count
        if cumulative >= half:
            return size_pt
    return samples[-1][0]


def _detect_formatting_heading(para, body_font_size_pt: Optional[float]) -> bool:
    """Detect if a paragraph is a heading based on visual formatting heuristics."""
    text = para.text.strip()
    if not text:
        return False

    if len(text) > _MAX_FMT_HEADING_LEN:
        return False
    if _LIST_ITEM_CHARS.search(text):
        return False
    words = text.split()
    if words and words[0] in _SENTENCE_STARTERS:
        return False

    runs = para.runs
    if not runs:
        return False

    total_chars = sum(len(r.text) for r in runs)
    if total_chars == 0:
        return False

    # Heuristic 1: Bold majority
    bold_chars = sum(len(r.text) for r in runs if r.bold is True)
    bold_fraction = bold_chars / total_chars

    if bold_fraction > _BOLD_MAJORITY_THRESHOLD:
        if len(runs) > 1 and runs[0].bold is True and runs[-1].bold is not True:
            pass
        else:
            return True

    # Heuristic 2: Larger font size
    if body_font_size_pt is not None:
        size_chars: dict[float, int] = {}
        for run in runs:
            if run.font.size is not None and len(run.text.strip()) > 0:
                sz = run.font.size.pt
                size_chars[sz] = size_chars.get(sz, 0) + len(run.text)
        if size_chars:
            dominant_size = max(size_chars, key=size_chars.get)
            if dominant_size >= body_font_size_pt + _MIN_FONT_DELTA_PT:
                return True

    # Heuristic 3: ALL CAPS
    word_count = len(words)
    if text.isupper() and len(text) <= _MAX_ALLCAPS_LEN and word_count <= _MAX_ALLCAPS_WORDS:
        return True
    allcaps_chars = sum(len(r.text) for r in runs if r.font.all_caps is True)
    if allcaps_chars / total_chars > _BOLD_MAJORITY_THRESHOLD:
        if len(text) <= _MAX_ALLCAPS_LEN and word_count <= _MAX_ALLCAPS_WORDS:
            return True

    return False


def _detect_numbered_heading(text: str) -> int:
    """Detect if text is a numbered heading pattern. Returns level (1-3) or 0."""
    if len(text) > _MAX_HEADING_LENGTH:
        return 0
    if _LIST_ITEM_CHARS.search(text):
        return 0
    m = _NUMBER_PREFIX_RE.match(text)
    if m:
        rest = m.group(1)
        words = rest.split()
        if words and words[0] in _SENTENCE_STARTERS:
            return 0
    for pattern, level in _NUMBERED_HEADING_PATTERNS:
        if pattern.match(text):
            return level
    return 0


def _has_page_break(para) -> bool:
    """Check if a paragraph contains or is preceded by a page break."""
    pPr = para._element.find(f'{_WORD_NS}pPr')
    if pPr is not None:
        if pPr.find(f'{_WORD_NS}pageBreakBefore') is not None:
            return True

    for run in para._element.findall(f'{_WORD_NS}r'):
        for br in run.findall(f'{_WORD_NS}br'):
            if br.get(f'{_WORD_NS}type') == 'page':
                return True
    return False


def _build_page_map(doc) -> List[int]:
    """Build a list of page numbers indexed by non-empty paragraph position."""
    page_numbers = []
    current_page = 1
    found_any_break = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if _has_page_break(para):
                current_page += 1
                found_any_break = True
            continue

        if _has_page_break(para):
            current_page += 1
            found_any_break = True

        page_numbers.append(current_page)

    if not found_any_break:
        return []
    return page_numbers


def _estimate_page(paragraph_index: int, paragraphs_per_page: int = 40) -> int:
    """Estimate page number from paragraph index."""
    return (paragraph_index // paragraphs_per_page) + 1


class DOCXParser(BaseParser):
    """Parser for DOCX documents with section_ref support."""

    def parse(self, file_path: Path) -> ParseResult:
        """Parse DOCX document. Falls back to raw ZIP/XML extraction on failure."""
        try:
            doc = Document(file_path)
        except Exception as e:
            logger.warning(
                "python-docx rejected %s (%s), attempting ZIP/XML fallback",
                file_path.name, e,
            )
            return self._parse_zip_fallback(file_path)
        return self._parse_standard(file_path, doc)

    def _parse_zip_fallback(self, file_path: Path) -> ParseResult:
        """Extract text from a .docx ZIP archive when python-docx rejects it."""
        import zipfile
        import xml.etree.ElementTree as ET
        import tempfile

        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        try:
            with zipfile.ZipFile(file_path) as zf:
                xml_candidates = ["word/document.xml"]
                xml_candidates.extend(
                    n for n in zf.namelist()
                    if n.endswith(".xml") and n not in xml_candidates
                )

                paragraphs: list[str] = []
                for candidate in xml_candidates:
                    if candidate not in zf.namelist():
                        continue
                    try:
                        tree = ET.parse(zf.open(candidate))
                        root = tree.getroot()
                        for t_elem in root.iter(f"{{{ns}}}t"):
                            if t_elem.text:
                                paragraphs.append(t_elem.text)
                    except ET.ParseError:
                        continue

                    if paragraphs:
                        break

                if not paragraphs:
                    for name in zf.namelist():
                        if name.endswith(".xml"):
                            try:
                                tree = ET.parse(zf.open(name))
                                for elem in tree.iter():
                                    if elem.text and elem.text.strip():
                                        paragraphs.append(elem.text.strip())
                            except ET.ParseError:
                                continue

        except zipfile.BadZipFile:
            return ParseResult(
                filename=file_path.name,
                sections=[Section(
                    title="Document",
                    content="Unable to read file: not a valid DOCX/ZIP archive.",
                    section_ref="page-1",
                )],
                raw_text="",
                page_count=1,
                metadata={"parser": "docx-failed", "error": "BadZipFile"},
            )

        raw_text = "\n".join(paragraphs)
        if not raw_text.strip():
            return ParseResult(
                filename=file_path.name,
                sections=[Section(
                    title="Document",
                    content="No readable text found in document.",
                    section_ref="page-1",
                )],
                raw_text="",
                page_count=1,
                metadata={"parser": "docx-zip-fallback", "error": "no_text_extracted"},
            )

        # Use TextParser to detect sections (it handles section_ref)
        from .text_parser import TextParser
        tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False)
        tmp_file = Path(tmp.name)
        tmp.close()
        try:
            tmp_file.write_text(raw_text, encoding="utf-8")
            text_parser = TextParser()
            text_result = text_parser.parse(tmp_file)
            sections = text_result.sections
        finally:
            tmp_file.unlink(missing_ok=True)

        logger.info(
            "ZIP/XML fallback extracted %d paragraphs, %d sections from %s",
            len(paragraphs), len(sections), file_path.name,
        )

        return ParseResult(
            filename=file_path.name,
            sections=sections,
            raw_text=raw_text,
            page_count=1,
            metadata={"parser": "docx-zip-fallback", "paragraph_count": len(paragraphs)},
        )

    def _parse_standard(self, file_path: Path, doc) -> ParseResult:
        """Standard python-docx parsing path."""

        page_map = _build_page_map(doc)
        has_page_breaks = len(page_map) > 0

        # Extract tables and their positions
        table_positions = get_table_paragraph_positions(doc)
        table_markdowns = []
        for table in doc.tables:
            markdown = table_to_markdown(table)
            if markdown:
                table_markdowns.append(markdown)

        table_map = {}
        for idx, pos in enumerate(table_positions):
            if idx < len(table_markdowns):
                table_map[pos] = table_markdowns[idx]

        if table_markdowns:
            logger.info(
                "Extracted %d tables from document as Markdown",
                len(table_markdowns)
            )

        sections = []
        raw_text = []
        current_section_title = None
        current_section_content = []
        current_section_page = 1
        paragraph_index = 0
        section_index = 0

        for para in doc.paragraphs:
            if paragraph_index in table_map:
                table_md = table_map[paragraph_index]
                raw_text.append(table_md)
                current_section_content.append(table_md)

            text = para.text.strip()
            if not text:
                paragraph_index += 1
                continue

            raw_text.append(text)

            if has_page_breaks and paragraph_index < len(page_map):
                current_page = page_map[paragraph_index]
            else:
                current_page = _estimate_page(paragraph_index)

            is_heading = False
            if para.style and para.style.name and para.style.name.startswith('Heading'):
                is_heading = True

            if is_heading:
                if current_section_title:
                    ref = _make_section_ref(current_section_title, section_index)
                    sections.append(Section(
                        title=current_section_title,
                        content='\n'.join(current_section_content),
                        section_ref=ref,
                        page_start=current_section_page,
                        parent_ref=_make_parent_ref(ref),
                    ))
                    section_index += 1
                current_section_title = text
                current_section_content = []
                current_section_page = current_page
            else:
                current_section_content.append(text)

            paragraph_index += 1

        if paragraph_index in table_map:
            table_md = table_map[paragraph_index]
            raw_text.append(table_md)
            current_section_content.append(table_md)

        if current_section_title:
            ref = _make_section_ref(current_section_title, section_index)
            sections.append(Section(
                title=current_section_title,
                content='\n'.join(current_section_content),
                section_ref=ref,
                page_start=current_section_page,
                parent_ref=_make_parent_ref(ref),
            ))

        detection_method = "style"

        if not sections:
            sections = self._parse_with_formatting(doc, page_map)
            detection_method = "formatting" if sections else detection_method

        if not sections:
            sections = self._parse_with_patterns(raw_text, page_map)
            detection_method = "pattern" if sections else detection_method

        if not sections:
            sections.append(Section(
                title="Document",
                content='\n'.join(raw_text),
                section_ref="page-1",
                page_start=1,
            ))
            detection_method = "single"

        max_page = page_map[-1] if page_map else (
            _estimate_page(paragraph_index - 1) if paragraph_index else 1
        )

        return ParseResult(
            filename=file_path.name,
            sections=sections,
            raw_text='\n'.join(raw_text),
            page_count=max(len(doc.sections), max_page),
            metadata={
                "parser": "python-docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(table_markdowns),
                "section_detection": detection_method,
                "page_source": "explicit_breaks" if has_page_breaks else "estimated",
            }
        )

    def _parse_with_formatting(
        self, doc, page_map: List[int]
    ) -> List[Section]:
        """Detect sections by visual formatting heuristics."""
        body_font_size = _compute_body_font_size(doc)
        has_page_breaks = len(page_map) > 0

        sections = []
        current_section_title = None
        current_section_content = []
        current_section_page = 1
        paragraph_index = 0
        section_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if has_page_breaks and paragraph_index < len(page_map):
                page = page_map[paragraph_index]
            else:
                page = _estimate_page(paragraph_index)

            if _detect_formatting_heading(para, body_font_size):
                if current_section_title:
                    ref = _make_section_ref(current_section_title, section_index)
                    sections.append(Section(
                        title=current_section_title,
                        content='\n'.join(current_section_content),
                        section_ref=ref,
                        page_start=current_section_page,
                        parent_ref=_make_parent_ref(ref),
                    ))
                    section_index += 1
                current_section_title = text
                current_section_content = []
                current_section_page = page
            else:
                current_section_content.append(text)

            paragraph_index += 1

        if current_section_title:
            ref = _make_section_ref(current_section_title, section_index)
            sections.append(Section(
                title=current_section_title,
                content='\n'.join(current_section_content),
                section_ref=ref,
                page_start=current_section_page,
                parent_ref=_make_parent_ref(ref),
            ))

        if sections:
            logger.info(
                "Formatting-based section detection found %d sections "
                "(body_font_size=%.1fpt)",
                len(sections),
                body_font_size or 0.0,
            )

        return sections

    def _parse_with_patterns(
        self, paragraphs: List[str], page_map: List[int]
    ) -> List[Section]:
        """Detect sections by numbered heading patterns in paragraph text."""
        has_page_breaks = len(page_map) > 0
        sections = []
        current_section_title = None
        current_section_content = []
        current_section_page = 1
        section_index = 0

        for idx, text in enumerate(paragraphs):
            level = _detect_numbered_heading(text)

            if has_page_breaks and idx < len(page_map):
                page = page_map[idx]
            else:
                page = _estimate_page(idx)

            if level > 0:
                if current_section_title:
                    ref = _make_section_ref(current_section_title, section_index)
                    sections.append(Section(
                        title=current_section_title,
                        content='\n'.join(current_section_content),
                        section_ref=ref,
                        page_start=current_section_page,
                        parent_ref=_make_parent_ref(ref),
                    ))
                    section_index += 1
                current_section_title = text
                current_section_content = []
                current_section_page = page
            else:
                current_section_content.append(text)

        if current_section_title:
            ref = _make_section_ref(current_section_title, section_index)
            sections.append(Section(
                title=current_section_title,
                content='\n'.join(current_section_content),
                section_ref=ref,
                page_start=current_section_page,
                parent_ref=_make_parent_ref(ref),
            ))

        if sections:
            logger.info(
                "Pattern-based section detection found %d sections (no Heading styles)",
                len(sections),
            )

        return sections
